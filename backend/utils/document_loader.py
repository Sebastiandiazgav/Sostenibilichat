import os
from typing import List, Dict, Any
from core.config import settings

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from unstructured.partition.auto import partition
except ImportError:
    partition = None

class Document:
    """Simple document class to replace LangChain's Document"""
    def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
        self.page_content = page_content
        self.metadata = metadata or {}

class DocumentLoader:
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP

    def load_documents_from_directory(self, directory_path: str) -> List[Document]:
        """Load all documents from a directory recursively."""
        documents = []
        print(f"Scanning directory: {directory_path}")

        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                print(f"Processing file: {file_path}")
                try:
                    docs = self._load_single_file(file_path)
                    documents.extend(docs)
                    print(f"Loaded {len(docs)} chunks from {file}")
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")

        print(f"Total documents loaded: {len(documents)}")
        return documents

    def _load_single_file(self, file_path: str) -> List[Document]:
        """Load a single file based on its extension."""
        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            content = ""

            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            elif file_extension == '.pdf':
                if PdfReader:
                    reader = PdfReader(file_path)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                else:
                    content = f"PDF reader not available for: {os.path.basename(file_path)}"

            elif file_extension in ['.xlsx', '.xls']:
                if pd:
                    excel_data = pd.read_excel(file_path, sheet_name=None)
                    content = ""
                    for sheet_name, df in excel_data.items():
                        content += f"Sheet: {sheet_name}\n"
                        content += df.to_string() + "\n\n"
                else:
                    content = f"Excel reader not available for: {os.path.basename(file_path)}"

            elif file_extension == '.csv':
                if pd:
                    df = pd.read_csv(file_path)
                    content = df.to_string()
                else:
                    content = f"CSV reader not available for: {os.path.basename(file_path)}"

            elif file_extension == '.docx':
                if DocxDocument:
                    doc = DocxDocument(file_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                else:
                    content = f"Word reader not available for: {os.path.basename(file_path)}"

            elif file_extension in ['.pptx', '.ppt']:
                # Try unstructured for PowerPoint
                if partition:
                    try:
                        elements = partition(file_path)
                        content = "\n".join([str(element) for element in elements])
                    except Exception as e:
                        content = f"Error processing PowerPoint {file_path}: {e}"
                else:
                    content = f"PowerPoint reader not available for: {os.path.basename(file_path)}"

            else:
                # Try unstructured for other formats
                if partition:
                    try:
                        elements = partition(file_path)
                        content = "\n".join([str(element) for element in elements])
                    except Exception:
                        return []
                else:
                    return []

            if not content.strip():
                print(f"No content extracted from {file_path}")
                return []

            # Split into chunks
            chunks = self._split_text(content)
            return [
                Document(
                    page_content=chunk,
                    metadata={"source": file_path, "chunk": i, "file_type": file_extension}
                )
                for i, chunk in enumerate(chunks)
            ]

        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return []

    def _split_text(self, text: str) -> List[str]:
        """Simple text splitting."""
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            if end < len(text):
                # Find a good break point
                break_point = text.rfind(' ', start, end)
                if break_point == -1:
                    break_point = end
                chunk = text[start:break_point]
                start = break_point + 1
            else:
                chunk = text[start:]
                start = len(text)

            if chunk.strip():
                chunks.append(chunk.strip())

        return chunks

    def load_and_split_documents(self) -> List[Document]:
        """Load all documents from multiple directories and split them."""
        all_documents = []

        # Get DOCS_PATHS from settings, fallback to old DOCS_PATH for compatibility
        try:
            docs_paths = getattr(settings, 'DOCS_PATHS', None)
            if docs_paths is None:
                docs_paths = [getattr(settings, 'DOCS_PATH', 'docs')]
            print(f"Using docs_paths: {docs_paths}")
        except AttributeError as e:
            print(f"AttributeError accessing settings: {e}")
            docs_paths = ['docs']

        for docs_path in docs_paths:
            if os.path.exists(docs_path):
                print(f"Loading documents from: {docs_path}")
                documents = self.load_documents_from_directory(docs_path)
                all_documents.extend(documents)
            else:
                print(f"Path does not exist: {docs_path}")

        print(f"Total documents loaded from all paths: {len(all_documents)}")
        return all_documents